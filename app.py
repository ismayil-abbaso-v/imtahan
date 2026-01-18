import streamlit as st
import re
import random
import smtplib
from email.mime.text import MIMEText
from docx import Document
from io import BytesIO
from datetime import datetime, timedelta

# ================== EMAIL SETTINGS ==================
SENDER_EMAIL = "SENIN_GMAIL@gmail.com"
SENDER_PASSWORD = "GMAIL_APP_PASSWORD"
# ==================================================

st.set_page_config(page_title="İmtahan Hazırlayıcı", page_icon="📝")

# ================== OTP EMAIL FUNCTION ==================
def send_otp_email(receiver_email, otp_code):
    msg = MIMEText(f"""
Salam!

Sizin giriş kodunuz:
{otp_code}

Bu kodu heç kimlə paylaşmayın.

İmtahan Sistemi
""")

    msg["Subject"] = "İmtahan Sistemi - Giriş Kodu"
    msg["From"] = SENDER_EMAIL
    msg["To"] = receiver_email

    try:
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except:
        return False

# ================== LOGIN PAGE ==================
def login_page():
    st.title("🔐 Giriş / Qeydiyyat")

    if "otp" not in st.session_state:
        st.session_state.otp = None
    if "email" not in st.session_state:
        st.session_state.email = None
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False

    email = st.text_input("📧 Gmail daxil edin")

    if st.button("📨 Kod göndər"):
        if "@gmail.com" not in email:
            st.error("Zəhmət olmasa düzgün Gmail daxil edin!")
        else:
            otp = str(random.randint(100000, 999999))
            sent = send_otp_email(email, otp)

            if sent:
                st.session_state.otp = otp
                st.session_state.email = email
                st.success("✅ Kod email-ə göndərildi!")
            else:
                st.error("❌ Email göndərilə bilmədi")

    if st.session_state.otp:
        code = st.text_input("🔢 6 rəqəmli kodu daxil edin")

        if st.button("🚀 Daxil ol"):
            if code == st.session_state.otp:
                st.session_state.logged_in = True
                st.success("🎉 Uğurla daxil oldunuz!")
                st.rerun()
            else:
                st.error("❌ Kod yanlışdır")

# ================== LOGIN PROTECTION ==================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    login_page()
    st.stop()

# ================== ORIGINAL SYSTEM ==================

@st.cache_data
def parse_docx(file):
    doc = Document(file)
    question_blocks = []
    paragraphs = list(doc.paragraphs)
    i = 0

    option_pattern = re.compile(r"^\s*[A-Ea-e][\)\.\:\-\s]+(.*)")
    question_pattern = re.compile(r"^\s*(\d+)\s*[.)]\s*(.*)")

    def is_numbered_paragraph(para):
        return para._p.pPr is not None and para._p.pPr.numPr is not None

    while i < len(paragraphs):
        para = paragraphs[i]
        text = ''.join(run.text for run in para.runs).strip()
        if not text:
            i += 1
            continue

        q_match = question_pattern.match(text)
        if q_match or is_numbered_paragraph(para):
            question_text = q_match.group(2).strip() if q_match else text.strip()
            i += 1
            options = []

            while i < len(paragraphs):
                option_text = ''.join(run.text for run in paragraphs[i].runs).strip()
                if not option_text:
                    i += 1
                    continue
                if question_pattern.match(option_text):
                    break
                match = option_pattern.match(option_text)
                if match:
                    options.append(match.group(1).strip())
                    i += 1
                else:
                    if len(options) < 5:
                        options.append(option_text)
                        i += 1
                    else:
                        break

            if len(options) >= 2:
                question_blocks.append((question_text, options))
        else:
            i += 1

    return question_blocks

@st.cache_data
def parse_open_questions(file):
    doc = Document(file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    questions = []
    for p in paragraphs:
        p = re.sub(r"^\s*\d+\s*[.)]?\s*", "", p)
        if p:
            questions.append(p)
    return questions

def create_shuffled_docx_and_answers(questions):
    new_doc = Document()
    answer_key = []

    for idx, (question, options) in enumerate(questions, start=1):
        new_doc.add_paragraph(f"{idx}) {question}")
        correct_answer = options[0]
        shuffled_options = options[:]
        random.shuffle(shuffled_options)

        for j, option in enumerate(shuffled_options):
            letter = chr(ord('A') + j)
            new_doc.add_paragraph(f"{letter}) {option}")
            if option.strip() == correct_answer.strip():
                answer_key.append(f"{idx}) {letter}")

    return new_doc, answer_key

# ================== PAGE SYSTEM ==================
if "page" not in st.session_state:
    st.session_state.page = "home"

if st.session_state.page == "home":
    st.title("📝 Testləri Qarışdır və Biliklərini Yoxla!")
    st.markdown("Zəhmət olmasa bir rejim seçin:")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        if st.button("📝 Özünü imtahan et"):
            st.session_state.page = "exam"
            st.rerun()
    with col2:
        if st.button("🎲 Sualları Qarışdır"):
            st.session_state.page = "shuffle"
            st.rerun()
    with col3:
        if st.button("🎫 Bilet İmtahanı"):
            st.session_state.page = "ticket"
            st.rerun()
    with col4:
        if st.button("🧮 Bal Hesablaması"):
            st.session_state.page = "score_calc"
            st.rerun()
    with col5:
        if st.button("ℹ️ İstifadə Qaydaları"):
            st.session_state.page = "help"
            st.rerun()
else:
    st.sidebar.title("⚙️ Menyu")
    if st.sidebar.button("🚪 Çıxış"):
        st.session_state.logged_in = False
        st.session_state.otp = None
        st.session_state.email = None
        st.rerun()
